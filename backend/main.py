from typing import Annotated
import docx
from fastapi import FastAPI, File, HTTPException, UploadFile, Depends, status
from fastapi.responses import HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from starlette.requests import Request
from backend.models import User
from backend.config import name, password
from backend.database import c, db
import uvicorn


app = FastAPI()

# Настройка директорий для статических файлов и шаблонов
app.mount("/static", StaticFiles(directory="frontend/static"), name="static")
templates = Jinja2Templates(directory="frontend")



async def get_user(user: Annotated[User, Depends()]):
    if user.login == name and user.password == password:
        return True
    else:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Неверный логин или пароль")


@app.post('/enter')
async def user_input(user: User, request: Request):
    # Попробуем получить пользователя
    try:
        access_granted = await get_user(user)
        return {"access_granted": access_granted}
    except HTTPException as e:
        raise e  # Передаем исключение дальше для обработки в JavaScript

@app.get("/")
async def read_root(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.get("/success")
async def success_page(request: Request):
    return templates.TemplateResponse("success.html", {"request": request})

@app.post("/uploadfile/")
async def upload_file(file: UploadFile = File(...)):
    if file.filename.endswith('.docx'):
        content = """
        <html>
        <head>
            <style>
                body {
                    font-family: Arial, sans-serif;
                    margin: 20px;
                    padding: 20px;
                    background-color: #f4f4f9;
                    color: #333; 
                }
                h1 {
                    color: #5a5a5a;
                }
                p {
                    background-color: #fff;
                    padding: 10px;
                    border-radius: 5px;
                    box-shadow: 0 0 5px rgba(0, 0, 0, 0.1);
                    margin: 10px 0;
                }
                .container {
                    max-width: 800px;
                    margin: auto;
                }
            </style>
        </head>
        <body>
            <div class="container">
                <h1>Наши услуги</h1>
        """
        
        full_text = ''  # Инициализируем переменную перед циклом
        doc = docx.Document(file.file)
        
        # Объединяем текст из параграфов
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text.strip()
            if paragraph_text:  # Проверяем, что параграф не пуст
                full_text += paragraph_text + "\n"
        
        # Объединяем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:  # Проверяем, что ячейка не пустая
                        full_text += cell_text + "\n"
        
        # Здесь вставляем собранный текст в базу данных
        if full_text:  # Проверяем, что text не пуст
            c.execute("INSERT INTO text(text) VALUES (?)", (full_text,))
            db.commit()

        # Формирование content ответа
        for paragraph in doc.paragraphs:
            content += f"<p>{paragraph.text}</p>"

        content += """
                </div>
            </body>
        </html>
        """
        return HTMLResponse(content)
    
    return HTMLResponse("""<h1>Неверный формат файла</h1>""")

@app.get("/history")
async def history_page(request: Request):
    c.execute("SELECT text FROM text")  # Запрос на получение всех текстов из таблицы
    rows = c.fetchall()  # Извлечение всех строк из результата

    # Формируем контент для отображения
    content_txt = """
        <html>
        <head>
            <style>
                body {
                    font-family: Arial, sans-serif;
                    margin: 20px;
                    background-color: #f4f4f9;
                    color: #333;
                }
                h1 {
                    color: #5a5a5a;
                    text-align: center;
                }
                .container {
                    max-width: 800px;
                    margin: auto;
                    padding: 20px;
                    background-color: #fff;
                    border-radius: 10px;
                    box-shadow: 0 2px 10px rgba(0, 0, 0, 0.1);
                }
                .entry {
                    margin-bottom: 20px; 
                }
                .entry h2 {
                    font-size: 1.2em;
                    color: #5a5a5a;
                }
                .entry p {
                    margin: 5px 0;
                    background-color: #e9ecef;
                    padding: 10px;
                    border-radius: 5px;
                }
            </style>
        </head>
        <body>
            <div class="container">
                <h1>История загруженных документов</h1>
    """

    for row in rows:
        # Каждая запись будет добавлена в начало
        content_txt = f"""
            <div class="entry">
                <h2>Добавление:</h2>
                <p>{row[0].replace("\n", "<br>")
                }</p>  <!-- Заменяем символ переноса строки на <br> для HTML -->
            </div>
        """
    
    content_txt += """
            </div>
        </body>
        </html>
    """
    
    return HTMLResponse(content_txt)
